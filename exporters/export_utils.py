"""Export transcription results to TXT, DOCX, PDF, Markdown, and SRT."""

import io
import textwrap
from datetime import datetime
from typing import Any, Dict, List


def _fmt_time(seconds: float) -> str:
    if seconds is None:
        return "0:00:00"
    h, rem = divmod(int(seconds), 3600)
    m, s = divmod(rem, 60)
    return f"{h}:{m:02d}:{s:02d}"


def _fmt_srt_time(seconds: float) -> str:
    """Format a time in seconds as an SRT timestamp: HH:MM:SS,mmm."""
    if seconds is None or seconds < 0:
        seconds = 0.0
    total_ms = int(round(seconds * 1000))
    ms = total_ms % 1000
    total_s = total_ms // 1000
    h, rem = divmod(total_s, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _speaker(seg: Dict, speaker_map: Dict[str, str]) -> str:
    return speaker_map.get(seg["speaker"], seg["speaker"])


# Subtitle layout limits (broadcast convention: ~42 chars/line, 2 lines/cue).
SRT_MAX_LINE = 42
SRT_MAX_LINES = 2


def _split_into_cues(text: str) -> List[str]:
    """Wrap text to SRT_MAX_LINE and group lines into cues of at most
    SRT_MAX_LINES lines each, so no cue overflows the screen."""
    lines = textwrap.wrap(text, SRT_MAX_LINE) or [text]
    cues = []
    for i in range(0, len(lines), SRT_MAX_LINES):
        cues.append("\n".join(lines[i : i + SRT_MAX_LINES]))
    return cues


def _meta(t: Dict) -> str:
    return (
        f"Language: {t['language']}  |  Engine: {t['engine']}  |  "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )


def export_txt(transcription: Dict[str, Any], speaker_map: Dict[str, str]) -> bytes:
    lines = [
        f"TRANSCRIPT: {transcription['filename']}",
        _meta(transcription),
        "=" * 72,
        "",
    ]
    for seg in transcription.get("segments", []):
        ts = f"{_fmt_time(seg['start'])} → {_fmt_time(seg['end'])}"
        lines.append(f"[{_speaker(seg, speaker_map)}]  {ts}")
        lines.append(seg["text"])
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_md(transcription: Dict[str, Any], speaker_map: Dict[str, str]) -> bytes:
    lines = [
        f"# Transcript: {transcription['filename']}",
        "",
        _meta(transcription),
        "",
        "---",
        "",
    ]
    for seg in transcription.get("segments", []):
        ts = f"{_fmt_time(seg['start'])} → {_fmt_time(seg['end'])}"
        lines.append(f"**{_speaker(seg, speaker_map)}** `{ts}`")
        lines.append("")
        lines.append(seg["text"])
        lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_srt(transcription: Dict[str, Any], speaker_map: Dict[str, str]) -> bytes:
    """Export as SubRip (.srt) subtitles.

    Cues contain only the spoken text — no speaker labels — for clean
    subtitles. Long segments are split into shorter cues (max two lines of
    ~42 characters) so they fit on screen, with the segment's time divided
    across the resulting cues in proportion to their text length.
    """
    blocks = []
    index = 1
    for seg in transcription.get("segments", []):
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        seg_start = seg.get("start") or 0.0
        seg_end = seg.get("end")
        if seg_end is None or seg_end < seg_start:
            seg_end = seg_start
        cues = _split_into_cues(text)
        total = sum(len(c) for c in cues) or 1
        cursor = seg_start
        for i, cue in enumerate(cues):
            # Give the last cue the exact segment end to avoid rounding drift.
            if i == len(cues) - 1:
                cue_end = seg_end
            else:
                cue_end = cursor + (seg_end - seg_start) * (len(cue) / total)
            blocks.append(
                f"{index}\n{_fmt_srt_time(cursor)} --> {_fmt_srt_time(cue_end)}\n{cue}\n"
            )
            cursor = cue_end
            index += 1
    # Cues are separated by a blank line, per the SRT convention.
    return "\n".join(blocks).encode("utf-8")


def export_docx(transcription: Dict[str, Any], speaker_map: Dict[str, str]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading(f"Transcript: {transcription['filename']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_para = doc.add_paragraph()
    run = meta_para.add_run(_meta(transcription))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    for seg in transcription.get("segments", []):
        ts = f"{_fmt_time(seg['start'])} → {_fmt_time(seg['end'])}"
        sp = _speaker(seg, speaker_map)

        header = doc.add_paragraph()
        r_name = header.add_run(sp)
        r_name.bold = True
        r_name.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
        r_ts = header.add_run(f"  {ts}")
        r_ts.font.size = Pt(8)
        r_ts.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_paragraph(seg["text"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_pdf(transcription: Dict[str, Any], speaker_map: Dict[str, str]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "TransTitle",
        parent=styles["Heading1"],
        fontSize=15,
        spaceAfter=4,
        alignment=TA_LEFT,
    )
    meta_style = ParagraphStyle(
        "TransMeta",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.gray,
        spaceAfter=10,
    )
    speaker_style = ParagraphStyle(
        "TransSpeaker",
        parent=styles["Normal"],
        fontSize=9,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#0066CC"),
        spaceBefore=10,
        spaceAfter=2,
    )
    text_style = ParagraphStyle(
        "TransText",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=4,
    )

    def _safe(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = [
        Paragraph(_safe(f"Transcript: {transcription['filename']}"), title_style),
        Paragraph(_safe(_meta(transcription)), meta_style),
        HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey),
        Spacer(1, 10),
    ]

    for seg in transcription.get("segments", []):
        ts = f"{_fmt_time(seg['start'])} → {_fmt_time(seg['end'])}"
        sp = _safe(_speaker(seg, speaker_map))
        story.append(
            Paragraph(
                f'{sp}  <font color="grey" size="7">{ts}</font>',
                speaker_style,
            )
        )
        story.append(Paragraph(_safe(seg["text"]), text_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── Summary exporters ──────────────────────────────────────────────────────

def export_summary_txt(transcription: Dict[str, Any]) -> bytes:
    summary = transcription.get("summary") or ""
    lines = [
        f"SUMMARY: {transcription['filename']}",
        _meta(transcription),
        "=" * 72,
        "",
        summary,
        "",
    ]
    return "\n".join(lines).encode("utf-8")


def export_summary_md(transcription: Dict[str, Any]) -> bytes:
    summary = transcription.get("summary") or ""
    lines = [
        f"# Summary: {transcription['filename']}",
        "",
        _meta(transcription),
        "",
        "---",
        "",
        summary,
        "",
    ]
    return "\n".join(lines).encode("utf-8")


def export_summary_docx(transcription: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    summary = transcription.get("summary") or ""

    doc = Document()
    title = doc.add_heading(f"Summary: {transcription['filename']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_para = doc.add_paragraph()
    run = meta_para.add_run(_meta(transcription))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    for line in summary.split("\n"):
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_summary_pdf(transcription: Dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    summary = transcription.get("summary") or ""

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "SumTitle", parent=styles["Heading1"], fontSize=15, spaceAfter=4, alignment=TA_LEFT
    )
    meta_style = ParagraphStyle(
        "SumMeta", parent=styles["Normal"], fontSize=8, textColor=colors.gray, spaceAfter=10
    )
    text_style = ParagraphStyle(
        "SumText", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=4
    )

    def _safe(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = [
        Paragraph(_safe(f"Summary: {transcription['filename']}"), title_style),
        Paragraph(_safe(_meta(transcription)), meta_style),
        HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey),
        Spacer(1, 10),
    ]
    for line in summary.split("\n"):
        if line.strip():
            story.append(Paragraph(_safe(line), text_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()
